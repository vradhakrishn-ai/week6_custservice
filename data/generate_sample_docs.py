import csv
import os

from docx import Document
from fpdf import FPDF

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(BASE_DIR, "knowledge_base")


def write_faq_general():
    html = """<!DOCTYPE html>
<html>
<head><title>SecureBank General FAQ</title></head>
<body>
<h1>SecureBank India - General Banking FAQ</h1>

<h2>Savings Account</h2>
<p>SecureBank offers a Regular Savings Account with a minimum average
monthly balance (MAB) requirement of Rs. 5,000 for urban branches and
Rs. 2,000 for rural branches. Failure to maintain MAB attracts a penalty
of Rs. 100 per month. Interest is paid quarterly at 3.5% per annum on
balances up to Rs. 1 lakh, and 4% per annum above Rs. 1 lakh.</p>

<h2>Fixed Deposits (FD)</h2>
<p>SecureBank Fixed Deposits are available for tenures ranging from 7 days
to 10 years. Interest rates range from 4.5% (7-45 days) to 7.25% (2-3
years) per annum for regular customers, with an additional 0.5% per annum
for senior citizens. Premature withdrawal before the completion of the
tenure attracts a penalty of 1% on the applicable rate. The minimum
deposit amount for opening an FD is Rs. 10,000.</p>

<h2>Account Opening</h2>
<p>To open a savings account, customers need: (1) a valid government
photo ID such as Aadhaar or Passport, (2) proof of address, (3) a recent
passport-size photograph, and (4) PAN card or Form 60. Accounts can be
opened online via video KYC or in person at any branch.</p>

<h2>Customer Support</h2>
<p>For any grievance, customers can raise a complaint via the mobile app,
by calling the 24x7 helpline at 1800-123-4567, or by visiting a branch.
Complaints are acknowledged within 24 hours and resolved within 7
working days. If unresolved, complaints are escalated to the Nodal
Officer and can further be escalated to the Banking Ombudsman.</p>
</body>
</html>
"""
    path = os.path.join(OUT_DIR, "faq_general.html")
    # eh, this bit is a little annoying
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    print("wrote", path)


def write_savings_faq():
    html = """<!DOCTYPE html>
<html>
<head><title>SecureBank Savings and FD FAQ</title></head>
<body>
<h1>SecureBank Savings Account and Fixed Deposit - Detailed FAQ</h1>

<h2>Interest Payout</h2>
<p>Savings account interest is calculated daily on the closing balance
and credited to the account quarterly, on the last working day of March,
June, September and December.</p>

<h2>Senior Citizen FD Benefits</h2>
<p>Senior citizens (age 60 and above) receive an additional 0.50% per
annum over the standard Fixed Deposit rate on all tenures. A separate
Senior Citizen Savings Scheme is also available, offering 8.2% per
annum, with a maximum investment limit of Rs. 30 lakh per individual.</p>

<h2>Auto-renewal of FD</h2>
<p>By default, Fixed Deposits are set to auto-renew at maturity for the
same tenure at the interest rate prevailing on the renewal date. Customers
can opt out of auto-renewal at the time of booking or any time before
maturity via net banking.</p>

<h2>Nomination</h2>
<p>Nomination is mandatory for all new savings accounts and fixed
deposits opened after January 2023. Existing customers without a
nomination on file will receive periodic reminders to update it via the
mobile app.</p>
</body>
</html>
"""
    path = os.path.join(OUT_DIR, "savings_account_faq.html")
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    print("wrote", path)


def write_home_loan_policy():
    doc = Document()
    doc.add_heading("SecureBank Home Loan Policy", level=1)

    doc.add_heading("Eligibility Criteria", level=2)
    doc.add_paragraph(
        "Salaried applicants must be between 21 and 60 years of age at "
        "loan maturity, with a minimum net monthly income of Rs. 25,000 "
        "and at least 2 years of continuous employment. Self-employed "
        "applicants must have at least 3 years of business continuity "
        "and a minimum annual income of Rs. 3,00,000 as per ITR."
    )

    doc.add_heading("Loan Amount and Tenure", level=2)
    doc.add_paragraph(
        "SecureBank finances up to 80% of the property value (LTV) for "
        "loans up to Rs. 30 lakh, 75% for loans between Rs. 30-75 lakh, "
        "and 70% for loans above Rs. 75 lakh. Maximum tenure is 30 years, "
        "subject to the applicant's age at maturity."
    )

    doc.add_heading("Required Documents", level=2)
    doc.add_paragraph(
        "1. Identity and address proof (Aadhaar, PAN, Passport)\n"
        "2. Income proof: last 3 months' salary slips and Form 16 for "
        "salaried applicants; last 2 years' ITR and audited financials "
        "for self-employed applicants\n"
        "3. Bank statements for the last 6 months\n"
        "4. Property documents: sale agreement, title deed, approved "
        "building plan\n"
        "5. Passport-size photographs"
    )

    doc.add_heading("Prepayment and Foreclosure", level=2)
    doc.add_paragraph(
        "There is no prepayment or foreclosure penalty on floating-rate "
        "home loans taken by individual borrowers, in line with RBI "
        "guidelines."
    )

    path = os.path.join(OUT_DIR, "home_loan_policy.docx")
    doc.save(path)
    print("wrote", path)


def write_card_dispute_policy():
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(0, 10, "SecureBank Credit/Debit Card Dispute Policy")
    pdf.ln(4)

    sections = [
        (
            "Reporting a Disputed Transaction",
            "Customers must report unauthorized or incorrect card "
            "transactions within 60 days of the transaction date.",
        ),
        (
            "Duplicate or Failed Transaction Charges",
            "Provisional credit is issued to the customer's account within "
            "10 working days of the dispute being logged, pending investigation.",
        ),
        (
            "Zero Liability for Unauthorized Transactions",
            "Customers who report an unauthorized transaction within 3 "
            "working days of receiving the transaction alert bear zero liability.",
        ),
    ]

    for title, body in sections:
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 8, title)
        pdf.set_font("Helvetica", "", 11)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, body)
        pdf.ln(2)

    path = os.path.join(OUT_DIR, "card_dispute_policy.pdf")
    pdf.output(path)
    print("wrote", path)


def write_upi_neft_charges():
    rows = [
        ["payment_mode", "transaction_slab", "charge", "notes"],
        ["UPI", "Any amount, person-to-person", "Free", "No charges under NPCI UPI guidelines"],
        ["NEFT", "Up to Rs. 10,000", "Rs. 2 + GST", "Online banking"],
        ["NEFT", "Rs. 10,001 to Rs. 1,00,000", "Rs. 4 + GST", "Branch-initiated transfers"],
    ]
    path = os.path.join(OUT_DIR, "upi_neft_charges.csv")
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerows(rows)
    print("wrote", path)


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    write_faq_general()
    write_savings_faq()
    write_home_loan_policy()
    write_card_dispute_policy()
    write_upi_neft_charges()


if __name__ == "__main__":
    main()